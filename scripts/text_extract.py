"""Shared document text extraction (Pass C / INFRA-027).

One place that turns a file path into plain text for the whole pipeline, so the
corpus loader, embedding store, convention parser, date cascade, and adaptive
spawn all accept the SAME input-format family:

    .pdf  .docx  .html/.htm  .txt  .md  .rst  .log  .json

Heavy/optional libraries (pypdf, python-docx, beautifulsoup4) are imported
lazily: a missing optional library degrades gracefully with a clear one-line
warning and an empty string, never a crash. Unsupported extensions are a
distinct signal (`is_supported` / `UnsupportedFormatError`) so every caller
WARNs and skips rather than silently dropping a file.
"""

from __future__ import annotations

import sys
from pathlib import Path

PLAINTEXT_EXTENSIONS = (".txt", ".md", ".rst", ".log", ".json")
PDF_EXTENSIONS = (".pdf",)
DOCX_EXTENSIONS = (".docx",)
HTML_EXTENSIONS = (".html", ".htm")
SUPPORTED_EXTENSIONS = (
    PLAINTEXT_EXTENSIONS + PDF_EXTENSIONS + DOCX_EXTENSIONS + HTML_EXTENSIONS
)


class UnsupportedFormatError(ValueError):
    """Raised by extract_text/extract_pages for an unrecognized extension."""


def is_supported(path) -> bool:
    return Path(path).suffix.lower() in SUPPORTED_EXTENSIONS


def warn_unsupported(path, *, where: str = "input") -> None:
    p = Path(path)
    print(
        f"[text_extract] WARN: unsupported file type "
        f"{(p.suffix.lower() or '(none)')!r} for {p.name!r} in {where}; skipping. "
        f"Supported: {', '.join(SUPPORTED_EXTENSIONS)}",
        file=sys.stderr, flush=True,
    )


def _warn_missing_lib(lib: str, path) -> None:
    print(
        f"[text_extract] WARN: optional library {lib!r} is not installed; cannot "
        f"read {Path(path).name!r}. Install it (see requirements.txt) to enable "
        f"this format. Skipping for now.",
        file=sys.stderr, flush=True,
    )


def _warn_read_error(kind: str, path, err: Exception) -> None:
    print(f"[text_extract] WARN: could not read {kind} {Path(path).name!r}: {err}",
          file=sys.stderr, flush=True)


def _extract_pdf_pages(path):
    try:
        import pypdf
    except ImportError:
        _warn_missing_lib("pypdf", path)
        return []
    try:
        reader = pypdf.PdfReader(str(path))
        return [(i + 1, (pg.extract_text() or "")) for i, pg in enumerate(reader.pages)]
    except Exception as e:
        _warn_read_error("PDF", path, e)
        return []


def _extract_docx(path) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        _warn_missing_lib("python-docx", path)
        return ""
    try:
        document = docx.Document(str(path))
        parts = [p.text for p in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                parts.append("\t".join(cell.text for cell in row.cells))
        return "\n".join(parts)
    except Exception as e:
        _warn_read_error("DOCX", path, e)
        return ""


def _extract_html(path) -> str:
    try:
        raw = Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        _warn_read_error("HTML", path, e)
        return ""
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        _warn_missing_lib("beautifulsoup4", path)
        return ""
    try:
        soup = BeautifulSoup(raw, "html.parser")  # builtin parser; no lxml needed
        for tag in soup(["script", "style"]):
            tag.decompose()
        return soup.get_text(separator="\n")
    except Exception as e:
        _warn_read_error("HTML", path, e)
        return ""


def _extract_plaintext(path) -> str:
    try:
        return Path(path).read_text(encoding="utf-8", errors="replace")
    except Exception as e:
        _warn_read_error("text", path, e)
        return ""


def extract_text(path) -> str:
    """Return the file's plain text. Returns '' if empty/unreadable or an
    optional library is missing. Raises UnsupportedFormatError for an
    unrecognized extension (callers should pre-filter with is_supported)."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in PLAINTEXT_EXTENSIONS:
        return _extract_plaintext(p)
    if ext in PDF_EXTENSIONS:
        return "\n".join(t for _, t in _extract_pdf_pages(p))
    if ext in DOCX_EXTENSIONS:
        return _extract_docx(p)
    if ext in HTML_EXTENSIONS:
        return _extract_html(p)
    raise UnsupportedFormatError(ext or "(no extension)")


def extract_pages(path):
    """Return [(page_no, text), ...]. PDFs yield their real pages; every other
    supported format yields a single (1, full_text) entry. Empty/unreadable
    files yield []. Raises UnsupportedFormatError for an unrecognized extension."""
    p = Path(path)
    ext = p.suffix.lower()
    if ext in PDF_EXTENSIONS:
        return [(n, t) for n, t in _extract_pdf_pages(p) if t]
    if ext in SUPPORTED_EXTENSIONS:
        text = extract_text(p)
        return [(1, text)] if text.strip() else []
    raise UnsupportedFormatError(ext or "(no extension)")
