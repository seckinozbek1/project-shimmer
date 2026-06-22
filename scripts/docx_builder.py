"""Tracked-changes + comments docx generator (Part XVIII Section D, item 7 of operator spec).

Produces a Word document with:
  - Page size A4, 1" margins, Times New Roman 12pt body
  - Document title 16pt bold, section headers 13pt bold, key terms italic
  - No headers, no footers, no decorative elements, no colored text
  - Tracked-change revision marks (real w:ins / w:del XML, not simulated)
  - Margin comments via real w:commentRangeStart/End + w:commentReference + comments.xml part
  - Each comment carries the precision references: CONV-*, REF-*

Implementation notes:
  python-docx 1.2.0 ships first-class paragraph/run APIs but no high-level
  tracked-changes or comments APIs. We construct the revision and comment
  XML elements directly via OxmlElement and lxml, attach the comments part
  to the package manually, and let python-docx serialize the result.
"""

from __future__ import annotations

import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import CONTENT_TYPE, RELATIONSHIP_TYPE
from docx.opc.part import Part
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

import language_detect


# ---------- direction-aware output (language-agnostic; decided from content) ---

def _apply_para_direction(paragraph_or_pe, text: str) -> None:
    """Set right-to-left base direction on a paragraph when its text is RTL.
    Adds <w:bidi/> and right justification to the paragraph properties. LTR text
    is left untouched (Word default). Direction is decided from the content's own
    script — no hardcoded language assumption."""
    if language_detect.text_direction(text) != "rtl":
        return
    pe = getattr(paragraph_or_pe, "_p", paragraph_or_pe)
    pPr = pe.get_or_add_pPr()
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))
    jc = pPr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        pPr.append(jc)
    jc.set(qn("w:val"), "right")


def _apply_run_direction(r, text: str) -> None:
    """Mark a run (w:r element) RTL when its own text is RTL, and tag its bidi
    language from the dominant RTL script. Per-run granularity handles mixed
    content (e.g. an LTR case name quoted inside an RTL paragraph): the RTL runs
    get <w:rtl/>, the LTR runs do not, and Word lays them out correctly within a
    bidi paragraph."""
    if language_detect.text_direction(text) != "rtl":
        return
    rPr = r.find(qn("w:rPr"))
    if rPr is None:
        rPr = OxmlElement("w:rPr")
        r.insert(0, rPr)  # rPr must be the first child of w:r
    tag = language_detect.rtl_script_tag(text)
    if tag and rPr.find(qn("w:lang")) is None:
        lang = OxmlElement("w:lang")
        lang.set(qn("w:bidi"), "ar-SA" if tag == "ar" else "he-IL")
        rPr.append(lang)
    if rPr.find(qn("w:rtl")) is None:
        rPr.append(OxmlElement("w:rtl"))


_AUTHOR = "Shimmer / AMENDMENT_DRAFTER"
_INITIALS = "AD"
_COMMENTS_PART_NAME = "/word/comments.xml"
_COMMENTS_CONTENT_TYPE = (
    "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"
)
_COMMENTS_RELATIONSHIP_TYPE = (
    "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
)


def _now_iso() -> str:
    return datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


class CommentsPart(Part):
    """Carrier for the comments.xml part."""

    @classmethod
    def new(cls, package):
        partname = package.next_partname("/word/comments%d.xml")
        # Per ECMA-376, a single comments part is expected at /word/comments.xml.
        partname = "/word/comments.xml"
        element = OxmlElement("w:comments")
        from docx.oxml import xmlchemy
        xml = _serialize(element)
        return cls(partname, _COMMENTS_CONTENT_TYPE, xml, package)


def _serialize(element) -> bytes:
    from lxml import etree
    return etree.tostring(element, standalone=True, xml_declaration=True, encoding="UTF-8")


class AmendmentDocxBuilder:
    """Builds the tracked-changes + comments docx for one operational document."""

    def __init__(self, *, title: str, body_text: str, amendments: list[dict],
                 author: str = _AUTHOR, initials: str = _INITIALS):
        self.title = title
        self.body_text = body_text
        self.amendments = list(amendments or [])
        self.author = author
        self.initials = initials
        self.doc = Document()
        self._apply_format()
        self._comments_root = OxmlElement("w:comments")
        self._comments_part_attached = False
        self._next_id = 1

    # ---------- format -----------------------------------------------------------------------

    def _apply_format(self) -> None:
        # A4 page + 1-inch margins
        section = self.doc.sections[0]
        section.page_height = Inches(11.69)
        section.page_width = Inches(8.27)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        # No headers / no footers
        for header in (section.header, section.first_page_header, section.even_page_header):
            for p in list(header.paragraphs):
                p.clear()
        for footer in (section.footer, section.first_page_footer, section.even_page_footer):
            for p in list(footer.paragraphs):
                p.clear()
        # Default Normal style: Times New Roman 12pt
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        rPr = normal.element.get_or_add_rPr()
        rFonts = rPr.find(qn("w:rFonts"))
        if rFonts is None:
            rFonts = OxmlElement("w:rFonts")
            rPr.append(rFonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rFonts.set(qn(attr), "Times New Roman")

    def _add_title(self) -> None:
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(self.title)
        run.bold = True
        run.font.size = Pt(16)
        run.font.name = "Times New Roman"
        _apply_para_direction(p, self.title)
        _apply_run_direction(run._r, self.title)

    def _add_header(self, text: str) -> None:
        p = self.doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = "Times New Roman"
        _apply_para_direction(p, text)
        _apply_run_direction(run._r, text)

    # ---------- comments part --------------------------------------------------------------

    def _ensure_comments_part(self) -> None:
        if self._comments_part_attached:
            return
        from docx.opc.packuri import PackURI
        package = self.doc.part.package
        partname = PackURI(_COMMENTS_PART_NAME)
        # If the package already has a comments part, reuse it; otherwise create one.
        for rel in self.doc.part.rels.values():
            if rel.reltype == _COMMENTS_RELATIONSHIP_TYPE:
                self._comments_part_attached = True
                return
        blob = _serialize(self._comments_root)
        part = Part(partname, _COMMENTS_CONTENT_TYPE, blob, package)
        self.doc.part.relate_to(part, _COMMENTS_RELATIONSHIP_TYPE)
        self._comments_part_attached = True
        self._comments_part = part

    def _flush_comments(self) -> None:
        if not self._comments_part_attached:
            return
        self._comments_part._blob = _serialize(self._comments_root)

    def _new_comment(self, text: str) -> int:
        cid = self._next_id
        self._next_id += 1
        c = OxmlElement("w:comment")
        c.set(qn("w:id"), str(cid))
        c.set(qn("w:author"), self.author)
        c.set(qn("w:initials"), self.initials)
        c.set(qn("w:date"), _now_iso())
        for line in text.splitlines():
            p = OxmlElement("w:p")
            r = OxmlElement("w:r")
            t = OxmlElement("w:t")
            t.set(qn("xml:space"), "preserve")
            t.text = line
            r.append(t)
            p.append(r)
            c.append(p)
        self._comments_root.append(c)
        return cid

    # ---------- runs (tracked changes + comment refs) --------------------------------------

    def _make_run(self, text: str, *, italic: bool = False) -> "OxmlElement":
        r = OxmlElement("w:r")
        if italic:
            rPr = OxmlElement("w:rPr")
            i = OxmlElement("w:i")
            rPr.append(i)
            r.append(rPr)
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        _apply_run_direction(r, text)
        return r

    def _make_ins_run(self, text: str) -> "OxmlElement":
        ins = OxmlElement("w:ins")
        ins.set(qn("w:id"), str(self._next_id))
        ins.set(qn("w:author"), self.author)
        ins.set(qn("w:date"), _now_iso())
        self._next_id += 1
        r = OxmlElement("w:r")
        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        _apply_run_direction(r, text)
        ins.append(r)
        return ins

    def _make_del_run(self, text: str) -> "OxmlElement":
        d = OxmlElement("w:del")
        d.set(qn("w:id"), str(self._next_id))
        d.set(qn("w:author"), self.author)
        d.set(qn("w:date"), _now_iso())
        self._next_id += 1
        r = OxmlElement("w:r")
        t = OxmlElement("w:delText")
        t.set(qn("xml:space"), "preserve")
        t.text = text
        r.append(t)
        _apply_run_direction(r, text)
        d.append(r)
        return d

    def _make_comment_range_start(self, cid: int) -> "OxmlElement":
        e = OxmlElement("w:commentRangeStart")
        e.set(qn("w:id"), str(cid))
        return e

    def _make_comment_range_end(self, cid: int) -> "OxmlElement":
        e = OxmlElement("w:commentRangeEnd")
        e.set(qn("w:id"), str(cid))
        return e

    def _make_comment_reference(self, cid: int) -> "OxmlElement":
        r = OxmlElement("w:r")
        ref = OxmlElement("w:commentReference")
        ref.set(qn("w:id"), str(cid))
        r.append(ref)
        return r

    # ---------- body assembly ---------------------------------------------------------------

    def _paragraph_with_amendment(self, paragraph_text: str, amendment: dict) -> None:
        """Render one body paragraph that contains a tracked change."""
        original = (amendment.get("original_text") or "").strip()
        proposed = amendment.get("proposed_text")
        action = amendment.get("action", "flag")
        comment_text = self._format_comment_body(amendment)
        cid = self._new_comment(comment_text)

        p = self.doc.add_paragraph()
        pe = p._p  # the w:p element
        # Base paragraph direction from the full paragraph text; individual runs
        # set their own direction (mixed LTR/RTL content is handled per run).
        _apply_para_direction(pe, paragraph_text)

        # Split paragraph_text around the original to anchor the revision marks
        if original and original in paragraph_text:
            before, after = paragraph_text.split(original, 1)
        else:
            before, after = paragraph_text, ""
            original = paragraph_text.strip() or original

        if before:
            pe.append(self._make_run(before))

        pe.append(self._make_comment_range_start(cid))

        if action == "rephrase" and proposed:
            pe.append(self._make_del_run(original))
            pe.append(self._make_ins_run(proposed))
        elif action == "reject":
            pe.append(self._make_del_run(original))
        else:  # flag | annotate — keep the text but anchor a comment + insertion of empty marker
            pe.append(self._make_run(original))

        pe.append(self._make_comment_range_end(cid))
        pe.append(self._make_comment_reference(cid))

        if after:
            pe.append(self._make_run(after))

    def _format_comment_body(self, amendment: dict) -> str:
        conv = amendment.get("convention_ref") or "CONV-???"
        loc = amendment.get("location") or "REF-????"
        action = amendment.get("action", "flag")
        severity = amendment.get("severity", "advisory")
        finding_type = amendment.get("finding_type", "—")
        body = (amendment.get("comment") or "").strip()
        context_refs = amendment.get("context_refs") or []
        # Genesis Part XXV: uncertain margin comments are explicitly
        # tagged so a reviewer skimming the docx never confuses them
        # with confident findings.
        is_uncertain = bool(amendment.get("uncertain"))
        header_prefix = "[UNCERTAIN] " if is_uncertain else ""
        if is_uncertain and not body.lstrip().startswith("[UNCERTAIN]"):
            body = "[UNCERTAIN] " + body
        lines = [
            f"{header_prefix}[{conv}] [{loc}] {finding_type} / {severity} / {action}",
            "",
            body,
        ]
        if context_refs:
            lines.append("")
            lines.append("Context references: " + ", ".join(context_refs))
        return "\n".join(lines)

    # ---------- public API ------------------------------------------------------------------

    def build(self, output_path: Path) -> Path:
        self._add_title()
        self._add_header("Source text with tracked changes")
        # Render body as paragraphs. If amendments target specific paragraphs, attach to those.
        paragraphs = _split_body(self.body_text)
        unattached = list(self.amendments)
        self._ensure_comments_part()
        for para_text in paragraphs:
            attached = None
            for a in unattached:
                ot = (a.get("original_text") or "").strip()
                if ot and ot in para_text:
                    attached = a
                    break
            if attached is not None:
                unattached.remove(attached)
                self._paragraph_with_amendment(para_text, attached)
            else:
                p = self.doc.add_paragraph()
                _apply_para_direction(p, para_text)
                if para_text:
                    p._p.append(self._make_run(para_text))
        # Any amendments that didn't anchor in the body get appended as standalone paragraphs.
        if unattached:
            self._add_header("Additional amendments (no anchor found in source body)")
            for a in unattached:
                anchor_text = (a.get("original_text") or "(unspecified location)").strip()
                self._paragraph_with_amendment(anchor_text, a)
        # Citation-references section (plain text reference for auditors)
        self._add_header("References cited by the amendments above")
        for i, a in enumerate(self.amendments, 1):
            p = self.doc.add_paragraph()
            p._p.append(self._make_run(
                f"Amendment {i}: convention {a.get('convention_ref', '?')}; "
                f"operational location {a.get('location', '?')}; "
                f"context refs {', '.join(a.get('context_refs') or []) or '(none)'}; "
                f"action {a.get('action', 'flag')}; severity {a.get('severity', 'advisory')}."
            ))
        self._flush_comments()
        output_path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(output_path))
        return output_path


_PARA_SPLIT_RE = re.compile(r"\n\s*\n+")


def _split_body(text: str) -> list[str]:
    paras = [p.strip() for p in _PARA_SPLIT_RE.split(text or "") if p.strip()]
    return paras or [""]


def build_amendments_docx(*, title: str, body_text: str, amendments: list[dict],
                          output_path: Path) -> Path:
    """Convenience wrapper."""
    builder = AmendmentDocxBuilder(title=title, body_text=body_text, amendments=amendments)
    return builder.build(output_path)
